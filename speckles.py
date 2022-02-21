import re
import os
import inquirer
import firebase_admin
import openpyxl as oxl
import rsa
from rsa import PrivateKey, PublicKey
from time import sleep
import time
from docx import Document
from alive_progress import alive_bar
from datetime import date as dt
from datetime import datetime
from colorama import Fore, Style
# from phaxio import PhaxioApi

# Firebase and auto_suggestion
from firebase_admin import credentials, firestore, storage
from prompt_toolkit import prompt, HTML
from prompt_toolkit.completion import WordCompleter, FuzzyCompleter

# Load public key from dir
pubKey = rsa.PublicKey.load_pkcs1(open("pubKey.pem", "rb").read())


# Search the doc(s) and replace data
def docx_replace_regex(doc_obj, regex, replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def decrypt(data):
    arr = []
    for char in data:
        temp = int(char) - 25
        arr.append(chr(temp))
        # print(char)

    return ''.join(arr)


def encrypt2(data):
    def split(word):
        return [char for char in word]

    arr = []
    chars = split(data)

    for char in chars:
        arr.append(ord(char) + 25)

    return arr


def main():
    os.system('cls')
    firebase_admin.initialize_app(credentials.Certificate(
        './sa.json'), {'storageBucket': 'fourpeaks-sc.appspot.com'})
    db = firestore.client()

    # Create Suggestions
    raw_docs = db.collection(u'Auto Suggestions').stream()
    docs = []
    suggestion_list = {'fax': [], 'phone': [], 'dr': [], 'procedure': [], 'surgeons': [
        'LaTowsky', 'Mai', 'Kaplan', 'Kundavaram', 'Stern', 'Klauschie', 'Schlaifer', 'Jones', 'Wong', 'Devakumar']}

    os.system('cls')

    # Makes the docs usable as dicts
    with alive_bar(len(docs), title='Creating suggestions', theme='classic') as bar:
        for doc in raw_docs:
            docs.append(doc.to_dict())
            # sleep(0.01)
        for doc in docs:
            # sleep(.25)
            for [key, value] in doc.items():
                # print(f'{key} ==> {value}')
                doc[key] = decrypt(value)
        keys = ['fax', 'phone', 'dr', 'procedure']

        for doc in docs:
            for key in keys:
                # Append suggestion to list if it doesn't already exitst
                try:
                    suggestion_list[key].index(doc[key])
                except ValueError:
                    suggestion_list[key].append(doc[key])
                    pass
            # sleep(0.5)
            bar()

    os.system('cls')
    import PyPDF2 as pdf

    doc_question = [inquirer.Checkbox('docs', message=f'What documents do you need? 📝 ', choices=[
                                      'Anticoagulant', 'A1c Tests', 'Pacemaker', 'Clearance'])]
    specialist_question = [inquirer.List(
        'specialist', message='Which specialist are you contacting? 🩺 ', choices=['Cardio', 'PCP', 'Endo', 'Pulmo', 'Onco', 'Nephro', 'Neuro'])]
    og_prompt = inquirer.prompt(doc_question)
    spec_prompt = inquirer.prompt(specialist_question)
    # print(og_prompt['docs'])

    raw_docs = []
    docs = [Document('./medrecs/faxcover.docx')]

    # Take choices and add them to raw_docs
    for choice in og_prompt['docs']:
        raw_docs.append(choice)

    for doc in raw_docs:
        docs.insert(0, Document(f'./medrecs/{doc}.docx'))

    og_prompt['docs'].append('Faxcover')
    patient = {
        "ptName": input(f'What is the {Fore.RED}name of the patient{Style.RESET_ALL}? 🧍\n'),
        "dateOfBirth": input(f"What is the patient's {Fore.RED}date of birth{Style.RESET_ALL}? 📅\n"),
        "procedureDate": input(f'What is the {Fore.RED}date of the procedure{Style.RESET_ALL}? 📅\n'),
        "procedureName": prompt(f'What is the procedure name? (i.e. PVP (Photovaporization of the prostate)) 🔪\n', bottom_toolbar=HTML(' 💭💤  Speckles is dreaming...   '), completer=FuzzyCompleter(WordCompleter(suggestion_list['procedure'])), complete_in_thread=True, complete_while_typing=True),
        "anesthesiologistName": prompt(f'Who is the surgeon? (Kaplan, Wong...) 🩺\n', bottom_toolbar=HTML('  😴 Speckles is waking up...   '), completer=FuzzyCompleter(WordCompleter(suggestion_list['surgeons'])), complete_in_thread=True, complete_while_typing=True),
        "drName": prompt(f'What is the name of the doctor you are contacting? (Name only! No "Dr." needed!)\n', bottom_toolbar=HTML('  👀 Speckles is watching...   '), completer=FuzzyCompleter(WordCompleter(suggestion_list['dr'])), complete_in_thread=True, complete_while_typing=True),
        "pNumber": prompt(f'What is the phone number of the facility you are faxing? 📞\n', bottom_toolbar=HTML('  🤔 Speckles is questioning life...   '), completer=FuzzyCompleter(WordCompleter(suggestion_list['phone'])), complete_in_thread=True, complete_while_typing=True),
        "fNumber": prompt(f'What is the number you are faxing to? 📠\n', bottom_toolbar=HTML('  ⚰️  Speckles is dead now...   '), completer=FuzzyCompleter(WordCompleter(suggestion_list['fax'])), complete_in_thread=True, complete_while_typing=True),
        "forms": og_prompt["docs"],
        "dateOfFax": dt.today().strftime("%B %d, %Y"),
        "numberOfPages": str(len(docs)),
        "urgency": 'URGENT',
        "yourName": 'Zackery H./Franchesca G.',
    }
    os.system('cls')

    # Input data into tables
    for doc in docs:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for word, replacement in patient.items():
                            word_re = re.compile(word)
                            docx_replace_regex(doc, word_re, replacement)

    # Input data into paragraphs
    for doc in docs:
        for word, replacement in patient.items():
            word_re = re.compile(word)
            docx_replace_regex(doc, word_re, replacement)

    # Save the changed files "as"
    x = 0
    with alive_bar(title='Applying information to templates', bar='fish') as bar:
        for doc in docs:
            doc.save(f"file_{x}.docx")
            # doc.save(f"pacemaker_{patient['ptName']}.docx")
            if docs.index(doc) != len(docs) - 1:
                x += 1
            bar()

    # Edit files to be PDFs
    print('Converting to PDFs...')
    os.system('powershell docx2pdf .')
    sleep(2)
    mergeFile = pdf.PdfFileMerger()
    pdfFiles = []
    while x >= 0:
        pdfFiles.append(pdf.PdfFileReader(f"file_{x}.pdf", 'rb'))
        x -= 1

    for pdf in pdfFiles:
        mergeFile.append(pdf)

    names = patient['ptName'].split(' ')
    os.system('powershell rm *.docx')
    os.system('powershell rm *.pdf')
    mergeFile.write(
        f'Request- {names[1]}, {names[0]} {spec_prompt["specialist"]} Med Recs Req.pdf')

    # Add to 'suggestion' collections
    new_info = {'fax': encrypt2(patient['fNumber']), 'phone': encrypt2(patient['pNumber']), 'drType': encrypt2(spec_prompt['specialist']),
                'dr': encrypt2(patient['drName']), 'procedure': encrypt2(patient['procedureName']), 'n': encrypt2(patient['ptName'])}
    db.collection('Auto Suggestions').document().set(new_info)

    def add_storage(pt):
        # time.time()
        stamp = str(time.time())
        encName = rsa.encrypt(pt['ptName'].encode(), pubKey)

        with open(f"{stamp}.bytes", "wb") as f:
            f.write(encName)
            f.close()

        bucket = storage.bucket()
        fileName = f'{stamp}.bytes'
        blob = bucket.blob(fileName)
        blob.upload_from_filename(fileName)
    add_storage(patient)
    # Write to excel
    wb = oxl.Workbook()  # Create workbook
    ws = wb.active  # Set active worksheet

    col = 1
    igs = ['urgency', 'yourName', 'numberOfPages', 'fNumber']
    patient['forms'] = ', '.join(patient['forms'])
    patient['pNumber'] += '/' + patient['fNumber']
    patient['pNumber'] = patient['pNumber'].replace(' ', '.')
    patient['pNumber'] = patient['pNumber'].replace('-', '.')
    patient['drName'] = f'{spec_prompt["specialist"]}: {patient["drName"]}'
    for key, value in patient.items():
        try:
            igs.index(key)  # See if key is ignored
        except ValueError:
            ws.cell(1, col, value)  # Write the data
            col += 1
        pass
    wb.save(
        f'Request- {names[1]}, {names[0]} {spec_prompt["specialist"]} Med Recs Req.xlsx')

    # Move patient to the scans folder

    print(f'Adding {patient["ptName"]}...')
    os.system(r"powershell mv *.pdf 'S:\'")
    print("Don't forget to add them to the log! Check the new '.xlsx' file!")
    sleep(5)
    os.system('cls')
    return True


main()

# def faxIt(pt):
#     fileFaxing = f"Request- {names[1]}, {names[0]} {spec_prompt['specialist']} Med Recs Req.pdf"
#     pt['fNumber'] = pt['fNumber'].replace(' ', '.')
#     pt['fNumber'] = pt['fNumber'].replace('-', '.')
#     faxNumber = pt['fNumber']
#     phaxio = PhaxioApi('7kgdmam2fdb4b3526751we7s5dim88s0u7n3kwxe',
#                        'bncxks0w76uods89k57h2kjvo8ykjxlytd306vnp')
#     phaxio.Fax.send(
#         to=faxNumber,
#         files=fileFaxing,
#         # direction='received'
#     )
# faxIt(patient)
